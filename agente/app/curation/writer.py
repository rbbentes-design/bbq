"""
Writer — etapa de escrita autoral do pipeline.

Calendário editorial fixo (Substack):
  Segunda  → week_ahead   : macro, dados econômicos da semana, pano de fundo
  Terça    → growth       : growth stocks, rotação, temas de crescimento
  Quarta   → flow_show    : fluxo, opções, CTAs, dealers, vol
  Quinta   → tese         : tese de ação específica ou tema macro
  Sexta    → week_recap   : recap da semana — fluxo, dados macro, parecer
  Sábado   → morning_call : fallback (pipeline não roda sábado normalmente)
  Domingo  → tese_livre   : tema aberto ao público — política, sociedade, tech

O modo é determinado pelo dia da semana do run_date.
O conteúdo e o ângulo emergem dos dados coletados.
"""

from __future__ import annotations

import json
import re
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Literal

from app.audit.logger import get_logger
from app.curation.llm_client import call_claude
from app.curation.models import CurationResult
from app.curation.persona import AUTHOR_PERSONA
from app.models.daily_ingestion_bundle import DailyIngestionBundle

_log = get_logger("curation.writer")

_MODEL = "claude-sonnet-4-6"

WritingMode = Literal[
    "week_ahead", "growth", "flow_show", "tese", "week_recap",
    "tese_livre", "morning_call", "podcast_sabado"
]

# ── Calendário editorial ────────────────────────────────────────────────────────

# weekday(): 0=Segunda, 1=Terça, 2=Quarta, 3=Quinta, 4=Sexta, 5=Sábado, 6=Domingo
_CALENDAR: dict[int, WritingMode] = {
    0: "week_ahead",
    1: "growth",
    2: "flow_show",
    3: "tese",
    4: "week_recap",
    5: "podcast_sabado",
    6: "tese_livre",
}


def _mode_for_date(run_date: str) -> WritingMode:
    try:
        d = date.fromisoformat(run_date)
    except ValueError:
        d = date.today()
    return _CALENDAR.get(d.weekday(), "morning_call")


# ── Domínios do brain por modo ─────────────────────────────────────────────────
# Cada mode puxa lenses apenas dos domínios relevantes.
# Tags disponíveis: macro, micro, trading, options, behavioral, narrative

_MODE_BRAIN_TAGS: dict[str, list[str]] = {
    # Editorial / público — todos os domínios relevantes (87 lenses de options
    # + 164 de trading informam o raciocínio mesmo em textos não-operacionais)
    "morning_call":   ["macro", "behavioral", "narrative", "options", "trading"],
    "week_ahead":     ["macro", "behavioral", "narrative", "options", "trading"],
    "week_recap":     ["macro", "behavioral", "narrative", "options", "trading"],
    "tese_livre":     ["macro", "behavioral", "narrative", "micro", "options", "trading"],
    "podcast_sabado": ["macro", "behavioral", "narrative", "micro", "options", "trading"],

    # Growth / tese de ação — micro + macro + behavioral + options/trading
    "growth":         ["macro", "micro", "behavioral", "options", "trading"],
    "tese":           ["macro", "micro", "behavioral", "options", "trading"],

    # Flow / operacional — options + trading + macro
    "flow_show":      ["options", "trading", "macro"],
}


# ── Instruções por modo ────────────────────────────────────────────────────────

_MODE_INSTRUCTIONS: dict[str, str] = {

    "week_ahead": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Week Ahead de segunda-feira. O leitor quer saber: o que está em jogo essa semana? \
Quais dados econômicos importam e por quê? Qual é o pano de fundo macro que vai \
colorir cada sessão? O que o mercado já está precificando e onde existe espaço para surpresa?

Use os dados da agenda econômica disponíveis no corpus e no RSS. Identifique os eventos \
que têm maior capacidade de mover mercado — não liste tudo, escolha o que tem peso real. \
Conecte a agenda com o posicionamento atual e com a narrativa macro de fundo.

Estrutura ideal: abertura forte com enquadramento da semana → 2-3 vetores principais com \
mecanismo → o que pode surpreender → como chegar preparado. Autoral, não calendário de banco. \
Dê uma leitura, não uma lista. 600-900 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Versão pública para fim do dia. Mais curta, mais afiada, mais compartilhável. \
Mostra inteligência suficiente para gerar desejo pelo conteúdo pago. Não entrega tudo — \
entrega o gancho certo. 200-350 palavras. Termina com uma linha que deixa o leitor \
querendo mais.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. 1 ideia forte por post. \
Impacto rápido, alta retenção, nada de frases vazias. Cada um precisa conter insight, \
ângulo ou provocação real. Separe claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Comunicação rápida, tom pessoal e direto. Fluida. Sem parecer institucional. \
Sensação de proximidade. 100-150 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Organize os pontos centrais da semana à frente em blocos objetivos:
AGENDA DA SEMANA: [releases e eventos com data]
DADOS DE MAIOR IMPACTO: [os que podem mover mercado]
NARRATIVA MACRO DE FUNDO: [1-2 frases]
POSICIONAMENTO ATUAL: [o que está precificado]
ONDE EXISTE SURPRESA: [upside/downside vs. consenso]
ATIVO/SETOR A MONITORAR: [1-2 destaques]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central do dia em 1-2 frases: a tensão, contradição ou implicação \
que diferencia este texto de um resumo genérico.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais que podem virar posts extras ou ser desenvolvidos em outro dia. \
Cada um em 1-2 frases (ÂNGULO A:, ÂNGULO B:).""",

    "growth": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Texto de terça sobre Growth Stocks. O foco é o segmento de crescimento: o que está \
comprimindo ou expandindo, por quê, e o que isso revela sobre o regime atual. \
Use dados do DeepVue Theme Tracker — Software, AI, Social Media, Cybersecurity, \
Biotech, Growth Stocks — comparando performance no dia, semana, mês e YTD.

A lógica central: taxa real explica múltiplo, múltiplo explica rotação, rotação \
revela o que o mercado acredita sobre crescimento e duration. Mostre essa cadeia. \
Identifique onde o consenso está errado sobre algum tema específico. Se houver nome \
de empresa ou setor com sinal claro no corpus ou no X, use. Não escreva como \
relatório setorial — escreva como quem entende a mecânica de precificação de \
crescimento. 600-900 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Versão pública para fim do dia. O insight central sobre crescimento em 200-350 palavras. \
Afiada, provocadora, compartilhável. Gera curiosidade pelo conteúdo pago sem entregar \
a análise completa.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. 1 ideia forte por post sobre \
growth, rotação, taxa real ou tema específico do dia. Nada vazio. \
Separe claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Tom pessoal e direto. O que está acontecendo com growth hoje e o que isso implica, \
em 100-150 palavras. Fluido, próximo, sem parecer institucional.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TEMAS QUE LIDERARAM: [com performance aproximada]
TEMAS QUE SOFRERAM: [com performance aproximada]
TAXA REAL / DURATION: [regime atual e implicação]
NARRATIVA DOMINANTE EM GROWTH: [em 1 frase]
SETOR/TEMA COM MELHOR ASSIMETRIA: [destaque do dia]
RISCO A MONITORAR: [o que pode mudar a leitura]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central em 1-2 frases.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais para posts extras (ÂNGULO A:, ÂNGULO B:).""",

    "flow_show": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Flow Show de quarta — voz CONVERSACIONAL E DIVERTIDA, como alguém explicando mercado \
em um bar para amigos inteligentes. NÃO é sell-side research. NÃO é PhD envergonhado. \
É o show semanal de opções/fluxo com humor, ironia seca e analogias terrenas.

REGRAS DE VOZ (obrigatórias neste modo):

1. ANALOGIAS TERRENAS no lugar de jargão puro. Exemplos do tom: "pombo jogando xadrez", \
"laranja sendo espremida", "pinball humano de bilhões", "café passado há três dias", \
"chegou depois que a pizza acabou", "aposta no empate e o time faz dois gols nos acréscimos". \
Se usar termo técnico, traga uma imagem visual junto.

2. TRADUTOR INLINE de jargão. Quando citar gamma, delta, 0DTE, HIRO, skew, GEX, gamma flip — \
explique em 1 frase curta no ato, com analogia. NÃO assuma que o leitor sabe. Ex: \
"Gamma é a velocidade com que o risco de uma opção muda. Perto do vencimento, fica instável \
como café passado há três dias." Faça isso com NATURALIDADE, não como glossário.

3. RITMO CONVERSACIONAL com pausas cômicas. Frases curtas alternadas com longas. Ironia seca. \
Ex: "Alguém ficou muito feliz. Outro, provavelmente, ficou sem dormir."

4. PULL QUOTES em blockquote com personalidade — frases compartilháveis: \
"O mercado não subiu porque as notícias melhoraram. Subiu porque quem apostou na queda foi \
forçado a comprar a alta. Poesia financeira." NÃO use "a convergência de fluxos suporta...".

5. DADO ESPECÍFICO + REAÇÃO HUMANA sempre juntos. "$1 para $17 em uma sessão. Alguém \
ficou muito feliz." NUNCA dado solto sem humanização.

6. PALAVRAS PROIBIDAS (tique de sell-side): "mecanismo de", "dissociação entre", \
"estrutura frágil", "convergência", "dinâmica de", "a tensão entre". Substitua por verbos \
concretos + analogia física.

7. FINAL QUE EMPODERA O LEITOR. Termine fazendo o leitor se sentir mais esperto que o \
comentarista de TV que explicou o rali como "otimismo macro". NÃO resuma como conclusão formal.

ESTRUTURA: Tudo conectado num único fio narrativo — a narrativa macro do dia aciona a \
mecânica de opções (gamma, dealers, 0DTE, skew), que se manifesta numa rotação setorial \
concreta, com um sinal de posicionamento por baixo que o consenso ainda não viu. Use os \
dados disponíveis: GEX por strike, gamma flip, call/put wall, dealer positioning, VVIX vs VIX, \
CTAs, vol control, rotação DeepVue. Parágrafos que carregam um ao outro, não seções separadas.

Abertura: gancho cotidiano que pega o leitor em 2 frases. Não "O mercado fechou em...". \
Feche com uma frase que amarra tudo e deixa o leitor se sentindo esperto. 700-1000 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Versão pública. O insight de fluxo mais impactante do dia em 200-350 palavras. \
Mostra que você enxerga o que a maioria não vê. Não entrega o mapa completo — \
entrega um ponto que deixa o leitor querendo o texto completo.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. Derivados do Flow Show. Podem ser sobre gamma, dealers, rotação, \
posicionamento sistemático, divergência narrativa/fluxo. 1 ideia técnica forte por post. \
Separe claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O que o fluxo está dizendo hoje, em 100-150 palavras. Tom direto e pessoal. \
Sem jargão excessivo. A leitura que o assinante precisa ter antes de abrir o mercado.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

DEALER POSITIONING: [regime atual — short/long gamma, nível de flip]
OPTIONS FLOW: [calls/puts dominantes, skew, volume relevante]
CTA / SISTEMÁTICO: [direção dos fluxos sistemáticos]
ROTAÇÃO SETORIAL: [o que está recebendo / perdendo fluxo]
DIVERGÊNCIA NARRATIVA vs. FLUXO: [se houver]
RISCO PARA AS PRÓXIMAS SESSÕES: [o que pode quebrar o regime]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central em 1-2 frases.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais para posts extras (ÂNGULO A:, ÂNGULO B:).""",

    "tese": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Tese de quinta-feira. O texto mais rigoroso da semana. Profundidade de analista CFA \
sênior com linguagem de estrategista autoral. Cada premissa construída, não declarada. \
Cada afirmação com mecanismo, dado ou lógica econômica. O leitor sai convencido.

ESTRUTURA OBRIGATÓRIA:
1. CONTEXTO E SETUP (2-3 parágrafos) — ambiente macro, forças em jogo, por que agora
2. A VISÃO DO CONSENSO (2 parágrafos) — o que o mercado precifica, por que parece razoável
3. ONDE O CONSENSO ERRA — O MECANISMO REAL (3-4 parágrafos) — a engrenagem ignorada, \
   com dados, relações causais, histórico de regimes similares
4. IMPLICAÇÕES SETORIAIS E DE ATIVOS (2-3 parágrafos) — valuation, múltiplos, duration
5. A ASSIMETRIA — POSICIONAMENTO (2 parágrafos) — o trade, upside/downside, convexidade
6. RISCOS E INVALIDADORES (1-2 parágrafos) — específicos, não genéricos
7. CONCLUSÃO — uma frase assinável que amarra tudo

TOM: analítico, autoral, denso mas fluido. Técnica embutida na narrativa. 1200-1800 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

A essência da tese em 250-400 palavras para o público gratuito. Mostra o problema \
e a contradição central. Não entrega o mecanismo completo nem o posicionamento — \
entrega o suficiente para despertar desejo pelo texto completo. Termina com gancho.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. Derivados da tese. Podem ser sobre o mecanismo, a contradição do \
consenso, a assimetria ou um dado específico que sustenta o argumento. 1 ideia forte \
por post. Separe claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

A tese em 100-150 palavras. Tom pessoal. "Hoje publiquei sobre X. A lógica é simples: \
[mecanismo em 2-3 frases]. O que o consenso está errando é [1 frase]. Link no bio." \
Direto e inteligente.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TESE CENTRAL: [em 1 frase]
PREMISSA 1: [dado/mecanismo]
PREMISSA 2: [dado/mecanismo]
PREMISSA 3: [dado/mecanismo]
ONDE O CONSENSO ERRA: [em 1-2 frases]
ATIVOS IMPLICADOS: [long/short/underweight — com razão]
HORIZONTE TEMPORAL: [dias/semanas/meses]
INVALIDADOR PRINCIPAL: [o que quebraria a tese]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central em 1-2 frases.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais para posts extras (ÂNGULO A:, ÂNGULO B:).""",

    "week_recap": """\
Você vai produzir CINCO entregáveis obrigatórios da sexta-feira. \
Separe cada um com o marcador exato indicado abaixo. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== WEEK RECAP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O Week Recap não é um diário de notícias. É um relatório curto e inteligente \
que explica a semana com foco em fluxo e macro.

PRIORIDADE ANALÍTICA (nesta ordem):
1. Fluxo — para onde o dinheiro foi
2. Dados econômicos que moveram expectativa ou preço
3. Eventos macro e políticos que definiram a semana
4. Comportamento de preço e posicionamento
5. Narrativa final da semana

O texto deve responder implicitamente: quem comprou, quem vendeu, onde houve \
rotação, quais dados alteraram expectativas, qual narrativa venceu no fechamento.

FLUXO como eixo central: mostre rotação (growth/value/small caps/defensivos/\
commodities/bonds/dólar/ouro/crédito), sinais de positioning sistemático \
(CTA, vol control, dealers, rebalanceamento), divergência entre narrativa e \
fluxo real quando existir.

DADOS ECONÔMICOS: para cada dado relevante que saiu na semana, diga o que saiu, \
por que importou, como o mercado leu, qual impacto em juros/dólar/ações/crédito. \
Não faça calendário descritivo. Faça análise.

PROGRESSÃO TEMPORAL: como a semana começou → o que mudou no meio → o que ganhou \
força → como terminou → leitura dominante no fechamento. Se houve mudança de \
narrativa, destaque.

FORMATO: texto corrido sem subtítulos mecânicos. Parágrafos curtos. Linguagem \
inteligente e natural, sem cara de IA. Sem floreio. Sem listas intermináveis. \
Sem reexplicar conceito óbvio para leitor avançado. Se ficar longo, consolide. \
Corte redundância. EXTENSÃO: 500-700 palavras máximo.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== COMENTÁRIO PARA POST ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Uma versão mais curta, forte e publicável. A principal mensagem da semana em \
150-250 palavras. Deve ter impacto imediato, linguagem afiada, e deixar evidente \
que existe profundidade por trás. É a vitrine — gera desejo pelo conteúdo completo. \
Pode terminar com uma pergunta ou gancho que provoque engajamento.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Crie de 3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. Derivados do report. Cada um com uma ideia específica \
forte que possa ser publicada ao longo do dia em X, Substack Notes ou WhatsApp. \
1 ideia por post. Impacto rápido. Alta retenção. Nada de frases vazias ou \
motivacionais. Cada peça precisa conter insight, ângulo, provocação ou leitura \
diferenciada. Separe os posts claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Artigo completo para publicação pública no fim do dia. Não é teaser nem resumo — \
é um artigo inteiro, com argumento desenvolvido, mecanismo, implicação e fechamento \
forte. Intercale imagens disponíveis ao longo do texto exatamente como no WEEK RECAP \
(use [IMAGEM: IMG-N | legenda]). Tom mais acessível que o conteúdo pago, mas sem \
perder sofisticação. O leitor gratuito lê um artigo de qualidade real. 500-700 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Organize a semana em blocos objetivos prontos para planilha, dashboard ou \
resumo executivo. Use este formato:

FLUXO POR REGIÃO: [EUA, Europa, Emergentes, outros]
FLUXO POR SETOR: [que setores ganharam / perderam fluxo]
FLUXO POR FATOR: [growth/value/defensivos/commodities/crédito/bonds/dólar/ouro]
DADOS ECONÔMICOS: [lista dos releases com número e leitura de mercado]
EVENTOS MACRO/POLÍTICOS: [os que moveram mercado]
ATIVOS QUE LIDERARAM: [com retorno aproximado se disponível]
ATIVOS QUE SOFRERAM: [com retorno aproximado se disponível]
NARRATIVA DOMINANTE: [em 1-2 frases]
RISCO/PONTO DE ATENÇÃO SEMANA SEGUINTE: [o que monitorar]""",

    "tese_livre": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Tese Livre de domingo — texto aberto ao público sobre um tema polêmico, controverso \
ou simplesmente interessante. Pode ser política, sociedade, tecnologia, economia \
comportamental, ou o que estiver em voga.

Maior alcance — vai para assinantes gratuitos. Precisa combinar acessibilidade com \
profundidade. O leitor não precisa entender mercado, mas precisa sair com uma ideia \
nova, perspectiva diferente, ou contradição que não havia percebido.

Estrutura: abertura provocativa que quebra expectativa → desenvolvimento com mecanismo \
e evidência → onde o consenso está errado → implicação inesperada → fechamento forte \
e memorável. Escreva como quem tem algo real a dizer e não tem medo de dizê-lo. \
800-1200 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Este modo já é público por natureza — entregue uma versão ainda mais condensada e \
afiada do argumento central, em 200-300 palavras. Ideal para compartilhar em redes \
e ampliar alcance.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. Derivados do tema. Provocações, contradições, insights específicos. \
Alta chance de compartilhamento. 1 ideia por post (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

100-150 palavras. Tom pessoal. "Escrevi hoje sobre [tema]. O ponto central é \
[1-2 frases]. Quem leu me disse que [reação esperada]. Link no bio." Direto e humano.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TEMA: [em 1 frase]
ARGUMENTO CENTRAL: [o mecanismo ou contradição central]
CONSENSO QUE DESMONTA: [o que a maioria acredita e está errado]
EVIDÊNCIA PRINCIPAL: [dado, fato ou lógica que sustenta]
IMPLICAÇÃO PARA O LEITOR: [por que isso importa concretamente]
GANCHO PARA SEMANA SEGUINTE: [como esse tema pode evoluir]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central em 1-2 frases.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais para posts extras (ÂNGULO A:, ÂNGULO B:).""",

    "podcast_sabado": """\
Você vai produzir o SCRIPT DE PODCAST de sábado. Separe cada bloco com o marcador exato. Não pule nenhum.

REGRA FUNDAMENTAL — O QUE É ESTE PODCAST:
Não é leitura de mercado. Não é recap de notícias. É um PODCAST DE IDEIAS POLÊMICAS.
Você escolhe UM tema específico e polêmico dos artigos do ZeroHedge Main, e desenvolve uma TESE sobre ele.
Uma tese tem começo, meio e fim. Tem uma pergunta central. Tem uma resposta que provoca.
O ouvinte deve terminar pensando: "Nunca tinha olhado por esse ângulo."

COMO ESCOLHER O TEMA:
Pegue o artigo mais polêmico, mais controverso, mais "o que está realmente acontecendo por trás disso" \
dos dados do ZeroHedge. Não o mais óbvio. O que tem mais camadas. O que conecta pontos que parecem \
não relacionados. Exemplos de boas perguntas: "Quem se beneficia com essa guerra?" \
"Por que esse escândalo foi ignorado pela grande mídia?" "O que esse movimento financeiro revela \
sobre quem realmente controla esse mercado?" "Essa crise é acidente ou design?"

COMO CONSTRUIR A NARRATIVA:
1. GANCHO (1 minuto): Uma pergunta ou fato que para o ouvinte imediatamente.
2. O ARTIGO / A REVELAÇÃO (2-3 minutos): Apresente o tema central — o que o ZeroHedge reportou, \
   o que está em jogo, quem são os atores.
3. O QUE ESTÁ POR TRÁS (4-5 minutos): Vá além do artigo. Conecte com outros fatos, padrões históricos, \
   incentivos. Desenvolva a tese. Use "e se..." e "isso não é coincidência porque...". \
   Deixe claro quando é especulação — "agora entra a parte que é minha leitura", \
   "isso ainda não está provado mas os indícios apontam para..." Seja intelectualmente honesto \
   mas corajoso.
4. IMPLICAÇÕES (2-3 minutos): O que isso significa para o ouvinte? Para o portfólio? Para o mundo real? \
   Para a narrativa que a mídia dominante está vendendo?
5. FECHAMENTO (1 minuto): A pergunta que fica. Algo que o ouvinte vai querer pesquisar depois.

TOM: Amigável, curioso, inteligente. Como alguém que descobriu algo interessante e quer \
compartilhar. Não panfletário. Não conspiracionista barato. Mas sem medo de ir onde a análise leva. \
Sem viés esquerda/direita — análise de estrutura, incentivos e consequências.

ESCRITA PARA VOZ — REGRAS OBRIGATÓRIAS:
- Frases curtas. Máximo 20 palavras por frase.
- Sem números escritos por extenso desnecessariamente. Use dígitos: "110 dólares", não "cento e dez dólares".
- Pausas naturais com reticências ou vírgulas. Ex: "E aí... vem a parte interessante."
- Sem jargão técnico sem explicação imediata.
- Sem listas com travessão. Flua como conversa.
- Use perguntas retóricas para criar ritmo: "Por quê? Porque..."

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== SCRIPT PODCAST ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Script completo para locução. Estrutura em 5 movimentos conforme acima. \
ESCREVA COM OPINIÃO E TESE — tome posição, desenvolva o argumento, provoque o ouvinte. \
Escreva exatamente como você falaria — ritmo natural, pausas marcadas, tom de conversa inteligente. \
Frases curtas. Números como dígitos. Sem markdown no meio do texto. 1.000-1.400 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TÍTULO DO EPISÓDIO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Um título forte para o episódio. Curto, provocativo, memorizável. \
Deve refletir o tema central — não a semana.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== DESCRIÇÃO DO EPISÓDIO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2-3 parágrafos para a descrição em plataformas de podcast (Spotify, Apple Podcasts). \
Apresenta o tema, provoca curiosidade, não entrega o argumento completo. 100-150 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== COMENTÁRIO DO POST ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Comentário fixado pelo autor no próprio post do Substack/X logo após publicar. \
Objetivo: abrir discussão, provocar resposta, gerar engajamento imediato. \
Tom pessoal, direto, como se estivesse falando com o leitor. \
Uma pergunta ou provocação que convida o ouvinte a responder. \
Máximo 3-4 linhas. Não repita o título nem a descrição — adicione um ângulo novo ou \
uma pergunta que ainda não foi respondida no episódio.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== SUBSTACK NOTES ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Um post único para o Substack Notes que cobre o episódio inteiro. \
Apresenta o tema, o argumento central, a provocação e convida para ouvir. \
Tom pessoal, fluido, como o autor escrevendo para seus assinantes. \
Não é resumo seco — é a voz do autor sobre o episódio completo. \
150-250 palavras. Termina com link implícito ou chamada clara para ouvir.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 posts curtos para X e WhatsApp. \
Cada um é uma frase ou ideia forte do episódio — gancho rápido, gera clique. \
Nada genérico. (POST 1:, POST 2:, etc.)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TEMA DO EPISÓDIO: [em 1 frase]
FONTE PRINCIPAL: [artigo ou tema do ZeroHedge]
ARGUMENTO CENTRAL: [o que o podcast defende ou questiona]
PROVOCAÇÃO INTELECTUAL: [a ideia que vai fazer o ouvinte pensar]
DURAÇÃO ESTIMADA: [em minutos]""",

    "morning_call": """\
Você vai produzir SETE entregáveis do dia. Separe cada um com o marcador exato. Não pule nenhum.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Morning Call. Objetivo e direto. Enquadre o dia, destaque os principais vetores, \
separe o que é motor do que é ruído. Não pareça clipping de manchetes. Ritmo forte, \
leitura rápida. O assinante lê antes de abrir o mercado e sai com uma leitura clara \
do que importa hoje. 400-600 palavras.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== TEXTO GRATUITO ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ponto mais relevante do dia em 150-250 palavras. Versão pública para fim do dia. \
Provoca curiosidade pelo conteúdo pago sem entregar a leitura completa.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== MICRO POSTS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

3 a 5 micro posts para X (Twitter). MAXIMO 280 CARACTERES CADA. 1 insight do dia por post. \
Nada vazio. Separe claramente (POST 1:, POST 2:, etc.).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== VERSÃO WHATSAPP ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Bom dia em 80-120 palavras. O que está em jogo hoje. Tom direto e pessoal.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== CONSOLIDAÇÃO ESTRUTURADA ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

VETOR PRINCIPAL DO DIA: [em 1 frase]
DADO/EVENTO QUE DEFINE O DIA: [o mais relevante]
ATIVO EM FOCO: [destaque do dia]
NARRATIVA DOMINANTE: [em 1 frase]
RISCO / PONTO DE ATENÇÃO: [o que pode surpreender]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULO PRINCIPAL ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

O ângulo editorial central em 1-2 frases.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
=== ÂNGULOS SECUNDÁRIOS ===
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

2 ângulos adicionais para posts extras (ÂNGULO A:, ÂNGULO B:).""",
}


# ── Contexto do corpus ─────────────────────────────────────────────────────────

def _build_curation_context(result: CurationResult) -> str:
    primary = result.narrative.primary_signal
    secondary = result.narrative.secondary_signals

    lines = [
        f"Data: {result.run_date}",
        f"Narrativa primária: {primary.label}",
        f"Confiança: {primary.confidence:.0%}",
        f"Descrição: {primary.description}",
        f"Verificação: {result.verification.overall_verdict}",
    ]
    if secondary:
        sec = secondary[0]
        if sec and sec.label and sec.label.upper() != "NONE":
            lines.append(f"Narrativa secundária: {sec.label} ({sec.confidence:.0%})")
            if sec.description:
                lines.append(f"Desc. secundária: {sec.description}")
    if primary.evidence_quotes:
        lines.append("\nEvidências-chave:")
        for q in primary.evidence_quotes[:5]:
            lines.append(f'  - "{q}"')
    if result.verification.hallucination_flags:
        lines.append(f"\nFlags: {', '.join(result.verification.hallucination_flags)}")
    return "\n".join(lines)


def _build_fred_context(fred_data: dict, mode: str) -> str:
    """Formata dados FRED para contexto do writer."""
    lines: list[str] = []

    # Agenda econômica da semana
    calendar = fred_data.get("calendar", [])
    # Filtra releases com impacto real
    _HIGH_IMPACT = [
        "employment situation", "nonfarm payroll", "consumer price", "cpi",
        "fomc", "federal open market", "gdp", "gross domestic", "jolts",
        "job openings", "retail sales", "pce", "personal consumption",
        "ism manufactur", "ism non-manufactur", "trade in goods",
        "unemployment insurance", "adp national", "case-shiller",
        "chicago fed", "chicago pmi", "financial conditions", "st. louis fed financial",
        "ppi", "producer price", "michigan consumer", "housing starts",
        "existing home", "treasury", "personal income",
    ]
    high_impact = [
        r for r in calendar
        if any(kw in r.get("release_name", "").lower() for kw in _HIGH_IMPACT)
    ]
    if high_impact:
        lines.append("=== AGENDA ECONÔMICA — PRÓXIMOS 10 DIAS (FRED) ===")
        seen = set()
        for r in high_impact:
            key = f"{r['date']}|{r['release_name']}"
            if key not in seen:
                seen.add(key)
                lines.append(f"  {r['date']}  {r['release_name']}")

    # Séries macro chave — últimos valores + variação
    series = fred_data.get("series", {})
    if series:
        lines.append("\n=== SÉRIES MACRO — FRED (últimos valores) ===")

        # Prioridade de categorias para o writer
        priority_cats = [
            "Política Monetária",
            "Inflação",
            "Mercado de Trabalho",
            "Crescimento",
            "Crédito e Condições Financeiras",
        ]
        for cat in priority_cats:
            cat_series = series.get(cat, [])
            if not cat_series:
                continue
            lines.append(f"\n{cat}:")
            for s in cat_series:
                val = s.get("value")
                chg = s.get("change")
                dt = s.get("date", "")
                if val is None:
                    continue
                chg_str = ""
                if chg is not None:
                    arrow = "▲" if chg > 0 else "▼" if chg < 0 else "─"
                    chg_str = f" {arrow}{abs(chg):.2f}"
                lines.append(
                    f"  {s['label']:40} {val:>10.2f} {s['unit']:12} [{dt}]{chg_str}"
                )

        # Default implícito do HY spread (se disponível)
        credit = series.get("Crédito e Condições Financeiras", [])
        hy = next((s for s in credit if s["series_id"] == "BAMLH0A0HYM2"), None)
        if hy and hy.get("value"):
            oas_bps = hy["value"] * 100  # converter de % para bps
            recovery = 0.30
            implied_dr = oas_bps / (1 - recovery)
            lines.append(
                f"\n  → HY OAS {oas_bps:.0f}bps implica taxa de default de "
                f"{implied_dr:.0f}bps ({implied_dr/100:.1f}% a.a.) assumindo recovery de 30¢"
            )

    return "\n".join(lines)


def _build_damodaran_context(damo: dict) -> str:
    """Formata dados Damodaran (ERP, country risk, WACC) para o writer."""
    lines: list[str] = []

    # ERP implícito — valor atual + histórico recente
    erp_current = damo.get("erp_current")
    erp_history = damo.get("erp_history", [])
    if erp_current:
        year = erp_current.get("year", "")
        erp  = erp_current.get("erp_pct")
        tbond = erp_current.get("t_bond_rate_pct")
        lines.append("=== DAMODARAN — EQUITY RISK PREMIUM (S&P 500) ===")
        tbond_str = f" | T-Bond: {tbond:.2f}%" if tbond else ""
        lines.append(f"  ERP Implícito atual ({year}): {erp:.2f}%{tbond_str}")

        # Últimos 5 anos para referência
        if len(erp_history) >= 5:
            recent = erp_history[-5:]
            hist_str = "  Histórico: " + " | ".join(
                f"{e['year']}: {e['erp_pct']:.1f}%" for e in recent
            )
            lines.append(hist_str)

        # Contexto: média histórica longa (todos os anos disponíveis)
        if len(erp_history) >= 10:
            avg = sum(e["erp_pct"] for e in erp_history if e.get("erp_pct")) / len(erp_history)
            lines.append(f"  Média histórica ({erp_history[0]['year']}–{erp_history[-1]['year']}): {avg:.2f}%")

    # Top setores por WACC (ordenados — mais alto = mais caro o capital)
    wacc_sectors = damo.get("wacc_sectors", [])
    if wacc_sectors:
        valid = [s for s in wacc_sectors if s.get("wacc_pct") is not None]
        sorted_sectors = sorted(valid, key=lambda s: s["wacc_pct"], reverse=True)

        lines.append("\n=== DAMODARAN — CUSTO DE CAPITAL POR SETOR (EUA) ===")
        lines.append("  [ Top 10 custo mais alto ]")
        for s in sorted_sectors[:10]:
            beta = f"  beta={s['beta']:.2f}" if s.get("beta") else ""
            lines.append(
                f"  {s['industry'][:45]:45}  WACC={s['wacc_pct']:.1f}%{beta}"
            )

        lines.append("  [ Top 10 custo mais baixo ]")
        for s in sorted_sectors[-10:][::-1]:
            beta = f"  beta={s['beta']:.2f}" if s.get("beta") else ""
            lines.append(
                f"  {s['industry'][:45]:45}  WACC={s['wacc_pct']:.1f}%{beta}"
            )

    # Country risk premiums — top 15 risco mais alto (mercados emergentes)
    country_risk = damo.get("country_risk", [])
    if country_risk:
        valid_cr = [c for c in country_risk if c.get("country_risk_premium_pct") is not None]
        sorted_cr = sorted(valid_cr, key=lambda c: c["country_risk_premium_pct"], reverse=True)
        lines.append("\n=== DAMODARAN — COUNTRY RISK PREMIUMS (top risco) ===")
        for c in sorted_cr[:15]:
            rating = f"  [{c['moody_rating']}]" if c.get("moody_rating") else ""
            crp = c["country_risk_premium_pct"]
            total = c.get("total_erp_pct", "")
            total_str = f"  total ERP={total:.1f}%" if total else ""
            lines.append(
                f"  {c['country']:30}  CRP={crp:.2f}%{total_str}{rating}"
            )

    return "\n".join(lines)


def _build_liquidity_context(liq: dict) -> str:
    """Formata indicadores de liquidez global para o writer. Todos os valores em USD."""
    lines: list[str] = ["=== LIQUIDEZ GLOBAL (todos os valores em USD bi) ==="]

    summary = liq.get("summary", {})
    us      = liq.get("us_liquidity", {})

    # ── Net Fed Liquidity ──────────────────────────────────────────────────────
    nfl = summary.get("net_fed_liquidity")
    if nfl:
        chg_str = f"  Δ1w={nfl['change_1w']:+.0f}" if nfl.get("change_1w") else ""
        lines.append(f"  Net Fed Liquidity         {nfl['value']:>8.0f} USD bi  [{nfl['date']}]{chg_str}")

    for sid, lbl in [("WALCL","  Fed Balance Sheet    "),
                     ("RRPONTSYD","  (-) RRP             "),
                     ("WTREGEN",  "  (-) TGA             ")]:
        e = us.get(sid)
        if e and e.get("value") is not None:
            chg_str = f"  Δ={e['change']:+.0f}" if e.get("change") else ""
            lines.append(f"{lbl}  {e['value']:>8.0f} USD bi  [{e['date']}]{chg_str}")

    # ── US Monetary aggregates ─────────────────────────────────────────────────
    for sid, lbl in [("M2SL","  US M2               "), ("BOGMBASE","  US Monetary Base    ")]:
        e = us.get(sid)
        if e and e.get("value") is not None:
            lines.append(f"{lbl}  {e['value']:>8.0f} USD bi  [{e['date']}]")

    # ── Money Market Funds ─────────────────────────────────────────────────────
    mmf_total = summary.get("money_market_total")
    mmf = liq.get("money_market", {})
    if mmf_total:
        lines.append(f"  Money Market Funds (tot)  {mmf_total['value']:>8.0f} USD bi  [{mmf_total['date']}]")
        for key, lbl in [("retail","    Retail"), ("institutional","    Institutional")]:
            entry = mmf.get(key)
            if entry:
                lines.append(f"{lbl:<26}  {entry['value']:>8.0f} USD bi")

    # ── Balanços Bancos Centrais Globais (em USD) ──────────────────────────────
    gbs = summary.get("global_balance_sheets")
    if gbs:
        lines.append(f"\n  Balanços BCx (G3 total)   {gbs['total_g3_usd_bi']:>8.0f} USD bi")
        for name, val in gbs.get("components", {}).items():
            lines.append(f"    {name:<22}  {val:>8.0f} USD bi")

    # ── M2 Global em USD ───────────────────────────────────────────────────────
    g_m2 = summary.get("global_m2_usd")
    if g_m2:
        lines.append(f"\n  M2 Global (G5)            {g_m2['total_g5_usd_bi']:>8.0f} USD bi")
        for country, val in g_m2.get("components", {}).items():
            if val:
                lines.append(f"    {country:<22}  {val:>8.0f} USD bi")
        fx = g_m2.get("fx_used", {})
        fx_parts = []
        if fx.get("EURUSD"):
            fx_parts.append(f"EUR/USD={fx['EURUSD']:.4f}")
        if fx.get("JPYUSD"):
            fx_parts.append(f"JPY/USD=1/{str(fx['JPYUSD']).replace('1/','')}")
        if fx.get("CNYUSD"):
            fx_parts.append(f"CNY/USD=1/{str(fx['CNYUSD']).replace('1/','')}")
        if fx_parts:
            lines.append("  [FX: " + "  ".join(fx_parts) + "]")

    # ── ECB M3 (se disponível) ─────────────────────────────────────────────────
    ecb_m3 = liq.get("ecb", {}).get("m3")
    if ecb_m3 and ecb_m3.get("value_usd_bi"):
        lines.append(f"\n  Euro Area M3              {ecb_m3['value_usd_bi']:>8.0f} USD bi  [{ecb_m3['date']}]")

    return "\n".join(lines)


def _build_bundle_context(bundle: DailyIngestionBundle | None, mode: str) -> str:
    """Extrai dados estruturados do bundle conforme o modo do dia."""
    if bundle is None:
        return ""

    sections: list[str] = []

    # SpotGamma — sempre relevante para flow_show, tese, week_recap
    if mode in ("flow_show", "tese", "week_recap", "morning_call"):
        for report in bundle.spotgamma_reports:
            if report.report_type == "FlowPatrol" and report.raw_text:
                sections.append(f"=== SpotGamma FlowPatrol ===\n{report.raw_text[:3000]}")
            elif report.report_type in ("PMNote", "FoundersNote") and report.raw_text:
                sections.append(
                    f"=== SpotGamma {report.report_type} — {report.title} ===\n{report.raw_text[:2000]}"
                )

    # DeepVue — sempre relevante (rotação de temas)
    deepvue_items = [r for r in bundle.rss_items if "DeepVue" in r.source_name]
    if deepvue_items:
        sections.append(f"=== DeepVue Theme Tracker ===\n{deepvue_items[0].summary[:2000]}")

    # Spectra — relevante especialmente na sexta
    if mode in ("week_recap", "tese_livre", "week_ahead"):
        spectra_items = [r for r in bundle.rss_items if "Spectra" in r.source_name]
        for item in spectra_items[:2]:
            if item.summary and len(item.summary) > 100:
                sections.append(f"=== Spectra Markets — {item.title} ===\n{item.summary[:1500]}")

    # FRED — agenda econômica + séries macro (week_ahead e week_recap)
    if mode in ("week_ahead", "week_recap") and bundle.fred_data:
        fred_ctx = _build_fred_context(bundle.fred_data, mode)
        if fred_ctx:
            sections.append(fred_ctx)

    # Damodaran — ERP, country risk, WACC por setor (week_ahead, tese, tese_livre)
    if mode in ("week_ahead", "tese", "tese_livre") and bundle.damodaran_data:
        damo_ctx = _build_damodaran_context(bundle.damodaran_data)
        if damo_ctx:
            sections.append(damo_ctx)

    # Liquidez Global — Net Fed Liquidity, MMF, ECB, M2 G4
    if mode in ("week_ahead", "week_recap", "flow_show", "tese") and bundle.global_liquidity:
        liq_ctx = _build_liquidity_context(bundle.global_liquidity)
        if liq_ctx:
            sections.append(liq_ctx)

    # Podcast de sábado: APENAS artigos do ZeroHedge principal — ignora tudo mais
    if mode == "podcast_sabado":
        zh_main = [r for r in bundle.rss_items if r.source_name == "ZeroHedge — Main"]
        if zh_main:
            podcast_sections: list[str] = []
            podcast_sections.append(
                "=== ARTIGOS DO ZEROHEDGE — ESCOLHA UM COMO TEMA CENTRAL ===\n"
                "Abaixo estão os artigos mais recentes e comentados do ZeroHedge.\n"
                "Escolha O MAIS POLÊMICO E COM MAIS CAMADAS para desenvolver a tese do episódio.\n"
                "NÃO use todos. Foque em UM. Vá fundo nele."
            )
            for item in zh_main:
                comments = ""
                for tag in (item.tags or []):
                    if tag.startswith("comments:"):
                        comments = f" [{tag.replace('comments:', '')} comentários]"
                podcast_sections.append(
                    f"TÍTULO: {item.title}{comments}\n"
                    f"URL: {item.url}\n"
                    f"CONTEÚDO:\n{item.summary[:3000]}"
                )
            return "\n\n".join(podcast_sections)
        else:
            # Sem ZH Main: usa RSS geral como fallback mas avisa
            fallback = ["=== ATENÇÃO: Sem artigos ZeroHedge Main. Use o RSS abaixo para escolher tema polêmico ==="]
            rss_all = [r for r in bundle.rss_items
                       if r.summary and len(r.summary) > 200
                       and "DeepVue" not in r.source_name]
            for item in rss_all[:8]:
                fallback.append(f"[{item.source_name}] {item.title}\n{item.summary[:800]}")
            return "\n\n".join(fallback)

    # RSS — filtrado por relevância para o modo
    rss_all = [r for r in bundle.rss_items
               if "DeepVue" not in r.source_name
               and "Spectra" not in r.source_name
               and r.summary and len(r.summary) > 200]

    if mode == "week_recap":
        # Para week_recap: separa dados econômicos do restante
        _MACRO_SOURCES = ("bls.gov", "bea.gov", "CalculatedRisk", "Calculated Risk",
                          "EconBrowser", "Federal Reserve", "Fed ")
        macro_data = [r for r in rss_all
                      if any(s.lower() in (r.source_name + r.title).lower()
                             for s in _MACRO_SOURCES)]
        other_rss = [r for r in rss_all if r not in macro_data]

        if macro_data:
            sections.append("=== RELEASES ECONÔMICAS DA SEMANA ===")
            for item in macro_data[:6]:
                sections.append(f"[{item.source_name}] {item.title}\n{item.summary[:600]}")

        if other_rss:
            sections.append("=== RSS / Corpus ===")
            for item in other_rss[:6]:
                sections.append(f"[{item.source_name}] {item.title}\n{item.summary[:400]}")
    else:
        max_rss = 5
        if rss_all:
            sections.append("=== RSS / Corpus ===")
            for item in rss_all[:max_rss]:
                sections.append(f"[{item.source_name}] {item.title}\n{item.summary[:500]}")

    # X Timeline — contas prioritárias relevantes para o modo
    if bundle.x_items:
        # LizAnn para week_recap
        if mode == "week_recap":
            lizann = [x for x in bundle.x_items if "LizAnn" in (x.author or "")]
            if lizann:
                sections.append("=== @LizAnnSonders ===")
                for t in lizann[:5]:
                    sections.append(f"  {t.text[:300]}")

        # Amostra geral do X para outros modos
        x_sample = bundle.x_items[:10]
        if x_sample:
            sections.append("=== X Timeline (amostra) ===")
            for t in x_sample:
                if t.text and len(t.text) > 30:
                    sections.append(f"  [{t.author}] {t.text[:200]}")

    return "\n\n".join(sections)


# ── Descoberta de tema (demand-driven) ────────────────────────────────────────

_THEME_SYSTEM = """\
Você é um editor-chefe de um newsletter de mercado com 10 mil leitores sofisticados. \
Seu trabalho NÃO é resumir dados. Seu trabalho é identificar QUAL É A CONVERSA DO DIA.

O que as pessoas estão falando? O que vão abrir o celular e querer entender? \
Qual é o tema com mais demanda de leitura HOJE?

MÉTODO:
1. Escaneie as manchetes RSS, timeline do X, e narrativa detectada.
2. Identifique o EVENTO ou TEMA que está dominando a atenção coletiva neste momento.
3. Pergunte: se eu fosse um investidor/gestor abrindo o celular HOJE, o que eu \
   quero saber? O que muda o meu trade de amanhã?
4. O tema precisa ter DEMANDA NATURAL de leitura, não ser apenas o dado mais \
   recente disponível.

REGRAS:
- O ingest é o UNIVERSO de possibilidades, não o tema. O tema vem da conversa pública.
- Nem tudo que está no ingest é relevante. Filtre pelo que GERA ENGAJAMENTO.
- Um tema forte tem: urgência (aconteceu agora), consequência (muda algo), \
  incerteza (o consenso não sabe a resposta), e conexão com dinheiro (afeta portfolio).
- DIA DA SEMANA importa: domingo = preparação para segunda. Sexta = recap. \
  Terça = rotação de growth. O tema precisa servir o leitor no contexto do dia.
- Se hoje é fim de semana e o mercado está fechado, NÃO invente ação de mercado. \
  Foque em: o que aconteceu que vai impactar a abertura de segunda?

Responda EXATAMENTE neste formato:

THEME: <o tema/evento dominante do dia — em 1 frase objetiva>
WHY_NOW: <por que HOJE? o que torna isso urgente agora?>
READER_QUESTION: <a pergunta que o leitor quer respondida — ex: "o que isso muda no meu trade de segunda?">
MARKET_CONNECTION: <como esse tema se conecta com portfolio/mercado — ex: "petróleo, treasuries, risk-off">
"""


def _discover_theme(result: CurationResult,
                    bundle: DailyIngestionBundle | None,
                    mode: str) -> dict[str, str]:
    """Identifica o tema do dia com base em demanda de leitura, não em oferta de dados.

    Retorna dict com keys: theme, why_now, reader_question, market_connection.
    """
    # Monta contexto com headlines, X timeline e narrativa detectada
    parts: list[str] = []

    parts.append(f"DATA: {result.run_date}")
    parts.append(f"DIA DA SEMANA: {_WEEKDAY_PT.get(mode, mode)}")
    parts.append(f"MODO EDITORIAL: {mode.upper()}")

    # Narrativa detectada pela curadoria
    primary = result.narrative.primary_signal
    parts.append(f"\nNARRATIVA DETECTADA: {primary.label}")
    parts.append(f"DESCRIÇÃO: {primary.description}")

    # Headlines RSS — só títulos para scanear rápido
    if bundle and bundle.rss_items:
        headlines = []
        for item in bundle.rss_items[:20]:
            if item.title and len(item.title) > 10:
                headlines.append(f"  [{item.source_name}] {item.title}")
        if headlines:
            parts.append("\nMANCHETES DO DIA (RSS):")
            parts.extend(headlines)

    # X Timeline — tweets mais recentes
    if bundle and bundle.x_items:
        tweets = []
        for t in bundle.x_items[:15]:
            if t.text and len(t.text) > 30:
                tweets.append(f"  [{t.author}] {t.text[:200]}")
        if tweets:
            parts.append("\nTIMELINE X:")
            parts.extend(tweets)

    # Evidências do corpus
    if primary.evidence_quotes:
        parts.append("\nEVIDÊNCIAS DO CORPUS:")
        for q in primary.evidence_quotes[:5]:
            parts.append(f'  - "{q}"')

    user_prompt = "\n".join(parts)

    raw = call_claude(_THEME_SYSTEM, user_prompt, model=_MODEL,
                      max_tokens=300, temperature=0.4)
    _log.info("writer_theme_discovered", mode=mode, raw=raw[:200])
    return _parse_theme(raw)


# Mapeamento de mode para dia da semana em PT (para contexto do LLM)
_WEEKDAY_PT: dict[str, str] = {
    "week_ahead":     "Segunda-feira",
    "growth":         "Terça-feira",
    "flow_show":      "Quarta-feira",
    "tese":           "Quinta-feira",
    "week_recap":     "Sexta-feira",
    "podcast_sabado": "Sábado",
    "tese_livre":     "Domingo",
    "morning_call":   "(dia variável)",
}


def _parse_theme(raw: str) -> dict[str, str]:
    result: dict[str, str] = {
        "theme": "", "why_now": "", "reader_question": "", "market_connection": "",
    }
    for line in raw.splitlines():
        for key, prefix in [
            ("theme", "THEME:"),
            ("why_now", "WHY_NOW:"),
            ("reader_question", "READER_QUESTION:"),
            ("market_connection", "MARKET_CONNECTION:"),
        ]:
            if line.strip().upper().startswith(prefix.upper()):
                result[key] = line.split(":", 1)[1].strip()
    return result


# ── Interpretação (ângulo + foco, dentro do tema) ─────────────────────────────

_ANGLE_SYSTEM = """\
Você é um estrategista de mercado com profundidade intelectual. O modo de escrita já \
está definido pelo calendário editorial. O TEMA DO DIA já foi definido pelo editor.

Sua tarefa: dado o tema do dia e os dados disponíveis, encontrar o MELHOR ÂNGULO \
para este texto. O ângulo não é o tema — é a LENTE específica que torna este texto \
diferente de qualquer outro sobre o mesmo assunto.

IMPORTANTE: O tema vem da demanda do leitor (o que as pessoas querem saber hoje). \
Os dados do ingest são EVIDÊNCIA para servir esse tema, não o tema em si. Use os dados \
para dar profundidade, números e mecanismo ao tema.

REGRA DE OURO: o ângulo NÃO pode ser apenas "a tensão entre A e B" ou "a divergência \
entre X e Y" ou "o paradoxo de Z". Esses são templates vazios que soam profundos mas \
não dizem nada. O ângulo precisa nomear o MECANISMO específico, a CONSEQUÊNCIA concreta, \
ou o ATOR que está sendo ignorado.

PENSAMENTO DE SEGUNDA ORDEM (obrigatório):
- O óbvio já está no preço e na manchete. Qual é a consequência INDIRETA?
- Qual estrutura, incentivo ou mecanismo invisível explica o que está acontecendo?
- Onde o consenso está errado, incompleto ou invertido?
- Se todo mundo está dizendo X, por que X pode estar errado?
- Qual é a conexão que ninguém está fazendo entre este tema e outro setor/ativo/dinâmica?

Responda EXATAMENTE neste formato:

FOCUS: <o mecanismo ou consequência específica que o texto deve expor — NÃO use "tensão entre", "divergência entre", "paradoxo de">
HOOK: <primeira frase ou ideia de abertura — algo concreto que prenda, não uma abstração>
ANGLE: <o ângulo específico que diferencia este texto — nomeie o mecanismo, o ator ou a consequência>
"""


def _find_angle(result: CurationResult, mode: str,
                bundle: DailyIngestionBundle | None = None,
                tema_hint: str | None = None) -> dict[str, str]:
    """Determina o ângulo, foco e hook — o modo já é fixo pelo calendário.

    Fluxo demand-driven:
    1. _discover_theme() identifica a CONVERSA do dia (o que o leitor quer saber)
    2. O tema guia a seleção de ângulo (dados do ingest são evidência, não tema)
    3. O ângulo conecta o tema com o trade/portfolio do leitor
    """

    # Podcast de sábado: ignora narrativa macro, usa só artigos ZH Main
    if mode == "podcast_sabado":
        bundle_ctx = _build_bundle_context(bundle, mode)
        if bundle_ctx:
            user_prompt = (
                "MODO DO DIA: PODCAST_SABADO\n\n"
                "Escolha UM artigo polêmico dos listados abaixo para ser o tema central do episódio.\n"
                "Ignore qualquer narrativa macro de mercado. Foque no tema com mais camadas e potencial de tese.\n\n"
                + bundle_ctx[:4000]
            )
            raw = call_claude(_ANGLE_SYSTEM, user_prompt, model=_MODEL, max_tokens=400, temperature=0.3)
            _log.info("writer_angle_done", mode=mode, raw=raw[:200])
            return _parse_angle(raw)

    # ── Step 1: Descobre o tema do dia (demand-driven) ────────────────────────
    theme = {"theme": "", "why_now": "", "reader_question": "", "market_connection": ""}
    if not tema_hint:
        try:
            theme = _discover_theme(result, bundle, mode)
            _log.info("writer_theme", theme=theme.get("theme", "")[:120])
        except Exception as exc:
            _log.warning("theme_discovery_failed", error=str(exc))

    # ── Step 2: Encontra o ângulo DENTRO do tema ──────────────────────────────
    curation_ctx = _build_curation_context(result)
    bundle_ctx = _build_bundle_context(bundle, mode)
    anti_rep = _build_anti_repetition_block(mode)

    # Brain: modelos mentais filtrados por domínio do mode
    brain_ctx = ""
    try:
        from app.curation.brain import get_brain_context
        narrative_label = result.narrative.primary_signal.label if result.narrative and result.narrative.primary_signal else ""
        brain_tags = _MODE_BRAIN_TAGS.get(mode, ["macro", "behavioral"])
        # Usa o tema descoberto como contexto de busca no brain (mais preciso)
        brain_search_ctx = theme.get("theme", "") or curation_ctx[:500]
        brain_ctx = get_brain_context(narrative_label, brain_search_ctx, use_llm=False, tags=brain_tags)
    except Exception as exc:
        _log.warning("brain_angle_error", error=str(exc))

    # Monta prompt com tema como driver principal
    user_prompt = f"MODO DO DIA: {mode.upper()}\n"
    user_prompt += f"DIA: {_WEEKDAY_PT.get(mode, mode)}\n\n"

    # Tema é o driver principal (se descoberto ou forçado)
    effective_theme = tema_hint or theme.get("theme", "")
    if effective_theme:
        user_prompt += f"═══ TEMA DO DIA (definido pelo editor) ═══\n"
        user_prompt += f"TEMA: {effective_theme}\n"
        if tema_hint:
            user_prompt += "(tema forçado pelo editor)\n"
        else:
            if theme.get("why_now"):
                user_prompt += f"POR QUE HOJE: {theme['why_now']}\n"
            if theme.get("reader_question"):
                user_prompt += f"PERGUNTA DO LEITOR: {theme['reader_question']}\n"
            if theme.get("market_connection"):
                user_prompt += f"CONEXÃO COM MERCADO: {theme['market_connection']}\n"
        user_prompt += (
            "\nO ângulo deve servir ESTE TEMA. Os dados abaixo são EVIDÊNCIA "
            "para dar profundidade, não são o tema em si.\n"
        )
    user_prompt += f"\n═══ DADOS DISPONÍVEIS (use como evidência) ═══\n{curation_ctx}"

    if brain_ctx:
        user_prompt += f"\n\n--- MODELOS MENTAIS ---\n{brain_ctx}"
    if anti_rep:
        user_prompt += f"\n\n{anti_rep}"
    if bundle_ctx:
        user_prompt += f"\n\n--- DADOS ESTRUTURADOS ---\n{bundle_ctx[:3000]}"

    raw = call_claude(_ANGLE_SYSTEM, user_prompt, model=_MODEL, max_tokens=400, temperature=0.5)
    _log.info("writer_angle_done", mode=mode, raw=raw[:200])
    return _parse_angle(raw)


def _parse_angle(raw: str) -> dict[str, str]:
    result: dict[str, str] = {"focus": "", "hook": "", "angle": ""}
    for line in raw.splitlines():
        if line.startswith("FOCUS:"):
            result["focus"] = line.split(":", 1)[1].strip()
        elif line.startswith("HOOK:"):
            result["hook"] = line.split(":", 1)[1].strip()
        elif line.startswith("ANGLE:"):
            result["angle"] = line.split(":", 1)[1].strip()
    return result


# ── Histórico de ângulos (anti-repetição) ──────────────────────────────────────

def _angle_log_path() -> Path:
    from app.storage.paths import workspace
    return workspace.state / "angle_history.json"


def _save_angle_log(mode: str, focus: str, angle: str, run_date: str, title: str = "") -> None:
    """Persiste o ângulo + título gerados para evitar repetição nas próximas execuções."""
    path = _angle_log_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        entries: list[dict] = json.loads(path.read_text(encoding="utf-8")) if path.exists() else []
    except Exception:
        entries = []
    entries.append({"date": run_date, "mode": mode, "focus": focus, "angle": angle, "title": title})
    # Mantém apenas últimas 60 entradas (2 meses)
    path.write_text(json.dumps(entries[-60:], ensure_ascii=False, indent=2), encoding="utf-8")


def _extract_title_from_text(text: str) -> str:
    """Extrai o primeiro título em negrito (**...**) do texto gerado."""
    import re as _re
    m = _re.search(r"\*\*([^*\n]{10,200})\*\*", text)
    return m.group(1).strip() if m else ""


def _load_angle_history(mode: str, days: int = 14) -> list[dict]:
    """Carrega ângulos recentes do mesmo modo (últimos N dias)."""
    path = _angle_log_path()
    if not path.exists():
        return []
    try:
        entries: list[dict] = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []
    today = date.today()
    cutoff = (today.toordinal() - days)
    result = []
    for e in entries:
        if e.get("mode") != mode:
            continue
        try:
            d = date.fromisoformat(e["date"]).toordinal()
            if d >= cutoff:
                result.append(e)
        except Exception:
            pass
    return result  # oldest first


def _build_anti_repetition_block(mode: str) -> str:
    """Retorna instrução de não-repetição baseada nos ângulos dos últimos 14 dias."""
    history = _load_angle_history(mode, days=14)
    if not history:
        return ""

    # Detecta patterns estruturais repetidos (templates vazios)
    _STRUCTURAL_PATTERNS = [
        "tensão entre", "divergência entre", "paradoxo de", "paradoxo do",
        "contradição entre", "desconexão entre", "dissociação entre",
        "a tensão", "a divergência", "o paradoxo",
    ]
    pattern_counts: dict[str, int] = {}
    for e in history:
        focus_lower = e.get("focus", "").lower()
        for pat in _STRUCTURAL_PATTERNS:
            if pat in focus_lower:
                pattern_counts[pat] = pattern_counts.get(pat, 0) + 1

    lines = [
        "\n⚠️  ÂNGULOS JÁ UTILIZADOS NOS ÚLTIMOS DIAS — NÃO REPITA NENHUM DESTES:",
        "Você DEVE escolher um ângulo completamente diferente dos listados abaixo.",
        "Se os dados do dia apontam para o mesmo tema, encontre uma faceta nova, "
        "uma contradição não explorada, uma consequência de segunda ordem ou um ativo/setor diferente.\n",
    ]
    for e in history[-7:]:  # mostra os últimos 7
        lines.append(f"  [{e['date']}] FOCUS: {e['focus'][:120]}")
        if e.get("angle"):
            lines.append(f"            ANGLE: {e['angle'][:120]}")
        if e.get("title"):
            lines.append(f"            TÍTULO: {e['title'][:120]}")

    # Lista exclusiva de títulos recentes (até 14 dias) — proibição forte
    recent_titles = [e.get("title", "").strip() for e in history if e.get("title")]
    recent_titles = [t for t in recent_titles if t][-10:]
    if recent_titles:
        lines.append("\n🚫 TÍTULOS JÁ USADOS — NÃO ESCREVA NADA PARECIDO COM ESTES:")
        for t in recent_titles:
            lines.append(f'  • "{t}"')
        lines.append(
            "Evite estruturas como 'O X que Y', 'O Rally/Cadeia/Mecanismo que ninguém precificou', "
            "'A semana em que...', 'Quando A encontra B'. Crie um título com VERBO ATIVO, "
            "nomeando o motor concreto (ex: 'Dealers vendem 8B em gamma na expiry de sexta')."
        )

    # Alerta de patterns estruturais
    if pattern_counts:
        lines.append("\n🚫 PATTERNS ESTRUTURAIS REPETIDOS — PROIBIDOS:")
        for pat, count in sorted(pattern_counts.items(), key=lambda x: -x[1]):
            if count >= 2:
                lines.append(f'  "{pat}" usado {count}x — NÃO USE esta construção.')
        lines.append(
            'Em vez de "tensão entre A e B", nomeie o MECANISMO: "dealers short gamma '
            'forçando amplificação de preço" ou "CTA crowding em 2 desvios criando '
            'squeeze mecânico". O ângulo é o motor, não o sintoma.'
        )

    lines.append(
        "\nA repetição de ângulo é o maior erro editorial possível. "
        "O leitor que voltou hoje quer VER o que mudou, não reler ontem com outras palavras."
    )
    return "\n".join(lines)


# ── Desk Intelligence context ─────────────────────────────────────────────────

def _build_wsb_context(bundle: DailyIngestionBundle | None) -> str:
    """Formata dados WSB + Squeeze para contexto do writer."""
    if bundle is None:
        return ""
    sw = bundle.swaggy_data if hasattr(bundle, "swaggy_data") else {}
    if not sw:
        return ""

    lines: list[str] = []

    wsb = sw.get("wsb_mentions", [])
    if wsb:
        lines.append("WSB TOP MENTIONS (WallStreetBets — sentimento varejo):")
        for t in wsb[:15]:
            ticker = t.get("ticker", "?")
            mentions = t.get("mentions", 0)
            rank = t.get("rank", 0)
            lines.append(f"  #{rank} {ticker:6s}  {mentions} menções")

    squeeze = sw.get("squeeze_candidates", [])
    if squeeze:
        lines.append("\nSQUEEZE CANDIDATES (alto short interest + catalisador):")
        for t in squeeze[:10]:
            ticker = t.get("ticker", "?")
            si = t.get("short_interest", t.get("si_pct", "?"))
            score = t.get("squeeze_score", t.get("score", "?"))
            lines.append(f"  {ticker:6s}  SI={si}  score={score}")

    return "\n".join(lines)


def _build_desk_intel_context(bundle: DailyIngestionBundle | None) -> str:
    """Carrega o DeskIntelligenceResult mais recente do dia e formata para o writer."""
    if bundle is None:
        return ""
    lines: list[str] = []

    # 1. SpotGamma Live data (do bundle)
    sg_live = getattr(bundle, "spotgamma_live", {}) or {}
    if sg_live and sg_live.get("tickers"):
        lines.append("SPOTGAMMA LIVE (Gamma/Dealers em tempo real):")
        spx = sg_live.get("spx_data")
        if spx:
            lines.append(f"  SPX: Gamma Flip={spx.get('gamma_flip', '?')}, "
                         f"Vol Trigger={spx.get('vol_trigger', '?')}, "
                         f"Dealer Regime={spx.get('dealer_regime', '?')}, "
                         f"Call Wall={spx.get('call_wall', '?')}, "
                         f"Put Wall={spx.get('put_wall', '?')}, "
                         f"GEX={spx.get('total_gex_b', '?')}B, "
                         f"Signal={spx.get('sg_signal', '?')}")
        for ticker, data in sg_live.get("tickers", {}).items():
            if ticker == "SPX":
                continue
            flip = data.get("gamma_flip", "")
            regime = data.get("dealer_regime", "")
            signal = data.get("sg_signal", 0)
            if flip or regime != "unknown":
                lines.append(f"  {ticker}: Flip={flip}, Regime={regime}, Signal={signal:.2f}")

    # 3. OptionsSnapshot (do options_store — dados BQuant)
    try:
        from app.providers.options_store import options_store as _opt_store
        snap = _opt_store.load_latest()
        if snap:
            lines.append(f"\nOPTIONS SNAPSHOT (BQuant Greeks Dashboard):")
            lines.append(f"  GEX Net: {snap.gex_net_bn:.2f}B | Gamma Regime: "
                         f"{'Long' if snap.gex_net_bn > 0 else 'Short'} Gamma")
            lines.append(f"  Gamma Flip: {snap.gamma_flip} | Squeeze Score: {snap.squeeze_score:.1f}")
            lines.append(f"  IV30d: {snap.iv_30d:.1f}% | RV30d: {snap.rv_30d:.1f}% | "
                         f"IV-RV Premium: {snap.iv_30d - snap.rv_30d:.1f}pp")
            lines.append(f"  Skew 25D: {snap.skew_25d:.2f} | P/C Ratio: {snap.pc_ratio:.2f}")
            lines.append(f"  Call Wall: {snap.call_wall} | Put Wall: {snap.put_wall}")
            lines.append(f"  Flow Score: {snap.flow_score_total:.2f} | "
                         f"Fragility: {snap.fragility:.2f} | Tail: {snap.tail_score:.2f}")
            lines.append(f"  Delta: {snap.delta_bn:.2f}B | Vanna: {snap.vanna_bn:.2f}B | "
                         f"Charm: {snap.charm_bn:.2f}B")
            if snap.vix:
                lines.append(f"  VIX: {snap.vix:.1f}")

            # Posicionamento CTA/VolCtrl/RP
            zs = snap.z_scores
            extreme = [(k, v) for k, v in zs.items() if abs(v) > 1.0]
            if extreme:
                extreme.sort(key=lambda x: abs(x[1]), reverse=True)
                lines.append(f"\n  POSICIONAMENTO INSTITUCIONAL (Z-scores):")
                for name, z in extreme[:8]:
                    direction = "LONG" if z > 0 else "SHORT"
                    intensity = "extremo" if abs(z) > 2 else "elevado"
                    lines.append(f"    {name}: Z={z:+.1f} ({direction} {intensity})")
    except Exception as exc:
        _log.debug("options_snapshot_context_skip", error=str(exc)[:60])

    if not lines:
        return ""
    return "\n".join(lines)


# ── Escrita ────────────────────────────────────────────────────────────────────

def _write(angle: dict[str, str], mode: str, result: CurationResult,
           bundle: DailyIngestionBundle | None = None) -> str:
    mode_instruction = _MODE_INSTRUCTIONS.get(mode, _MODE_INSTRUCTIONS["morning_call"])

    primary = result.narrative.primary_signal
    secondary = result.narrative.secondary_signals

    context_lines = [
        f"Data: {result.run_date}",
        f"Dia: {_WEEKDAY_PT.get(mode, mode)}",
    ]

    # Narrativa detectada como CONTEXTO (nunca como tema principal)
    context_lines.append(f"\nNarrativa de mercado detectada (contexto): {primary.label}")
    context_lines.append(f"Descrição: {primary.description}")
    if secondary:
        sec = secondary[0]
        if sec and sec.label and sec.label.upper() != "NONE":
            context_lines.append(f"Narrativa secundária: {sec.label} — {sec.description}")
    if primary.evidence_quotes:
        context_lines.append("\nEvidências do corpus:")
        for q in primary.evidence_quotes[:5]:
            context_lines.append(f'  - "{q}"')

    # Catálogo de imagens disponíveis
    catalog: list[dict] = []
    if bundle is not None:
        catalog = _build_image_catalog(bundle)

    context_lines.extend([
        f"\nFoco: {angle['focus']}",
        f"Ângulo: {angle['angle']}",
        f"Ideia de abertura: {angle['hook']}",
        f"\nMODO: {mode.upper()}",
        f"\nInstrução de formato:\n{mode_instruction}",
        "\n⚠️ TIMING DE DADOS — REGRA ABSOLUTA:",
        "  - Dados de performance diária (DeepVue 'Última sessão', market_prices 'daily_return')",
        "    referem-se à ÚLTIMA SESSÃO FECHADA, que pode ser ONTEM se o texto for escrito pre-market.",
        "  - NUNCA escreva 'Bitcoin subiu X% hoje' a menos que tenha certeza que a sessão de hoje já fechou.",
        "  - Use 'na última sessão', 'ontem', 'na sessão recente', 'no último pregão' quando houver",
        "    qualquer dúvida. NUNCA afirme 'hoje' cegamente a partir de um número diário.",
        "  - FIM DE SEMANA: se hoje é sábado ou domingo, o mercado está FECHADO. Não invente",
        "    movimentos de mercado. Foque no que aconteceu e no que IMPACTA a abertura de segunda.",
        "",
        "⚠️ PENSAMENTO ESTRUTURAL — OBRIGATÓRIO:",
        "  - Antes de aceitar qualquer explicação de superfície ('incerteza macro', 'budget freeze',",
        "    'risk-off'), pergunte: qual é o MECANISMO ESTRUTURAL por trás disso?",
        "  - Sempre conecte: FATO (o que aconteceu) → MECANISMO (por que/como) → IMPLICAÇÃO (e daí?)",
        "  - O leitor é trader. Cada insight precisa terminar em: 'e isso significa X para o seu trade'.",
        "  - Pense fora da caixa: que conexão entre setores/ativos/dinâmicas o consenso NÃO está fazendo?",
        "",
        "⚠️ ANTI-REDUNDÂNCIA — REGRA ABSOLUTA:",
        "  - Cada seção é uma entrega DIFERENTE. NUNCA reescreva o TEXTO PRINCIPAL em formato menor.",
        "  - TEXTO GRATUITO ≠ resumo do TEXTO PRINCIPAL. É um OUTRO ângulo da mesma ideia.",
        "  - MICRO POSTS = 3-5 ideias INDEPENDENTES, cada uma com seu próprio insight.",
        "  - VERSÃO WHATSAPP = tom pessoal, conversacional, nada reescrito.",
        "  - CONSOLIDAÇÃO = bullets estruturados de dados, não narrativa.",
        "  - Dentro do TEXTO PRINCIPAL: cada parágrafo deve AVANÇAR o argumento.",
        "    NÃO repita a mesma ideia reformulada. NÃO use a mesma imagem duas vezes.",
        "    NÃO encerre o texto repetindo o parágrafo de abertura.",
        "  - NUNCA use frases de impacto repetidas entre textos de dias diferentes.",
        "    Cada texto precisa de abertura e fechamento ÚNICOS.",
    ])

    # Briefing das imagens disponíveis (antes dos dados estruturados)
    briefing = _image_briefing(catalog)
    if briefing:
        context_lines.append(briefing)

    # Brain: modelos mentais filtrados por domínio do mode
    try:
        from app.curation.brain import get_brain_context
        narrative_label = primary.label if primary else ""
        brain_tags = _MODE_BRAIN_TAGS.get(mode, ["macro", "behavioral"])
        brain_block = get_brain_context(narrative_label, angle.get("focus", ""), use_llm=False, tags=brain_tags)
        if brain_block:
            context_lines.append(f"\n--- MODELOS MENTAIS (use como lentes, não cite fontes) ---\n{brain_block}")
    except Exception as exc:
        _log.warning("brain_write_error", error=str(exc))

    # Desk Intelligence: regime, opportunity, fragility, positioning
    desk_intel_ctx = _build_desk_intel_context(bundle)
    if desk_intel_ctx:
        context_lines.append(f"\n--- DESK INTELLIGENCE (dados quantitativos do pipeline) ---\n{desk_intel_ctx}")

    # WSB / Squeeze data
    wsb_ctx = _build_wsb_context(bundle)
    if wsb_ctx:
        context_lines.append(f"\n--- WSB + SQUEEZE (sentimento varejo / short squeeze) ---\n{wsb_ctx}")

    bundle_ctx = _build_bundle_context(bundle, mode)
    if bundle_ctx:
        context_lines.append(f"\n--- DADOS ESTRUTURADOS (SpotGamma / DeepVue / RSS / X) ---\n{bundle_ctx}")

    user_prompt = "\n".join(context_lines)

    # Todos os modos agora produzem múltiplos entregáveis — tokens expandidos
    max_tokens = 8000 if mode in ("tese", "tese_livre") else 6000
    text = call_claude(AUTHOR_PERSONA, user_prompt, model=_MODEL, max_tokens=max_tokens, temperature=0.7)
    _log.info("writer_text_done", mode=mode, chars=len(text))

    # Resolve IDs de imagem → sentinel <<<IMG:path>>>
    if catalog:
        text = _resolve_image_ids(text, catalog)
    return text


# ── Image catalog e embedding ──────────────────────────────────────────────────

def _build_image_catalog(bundle: DailyIngestionBundle) -> list[dict]:
    """Constrói catálogo de imagens com IDs sequenciais simples (IMG-1, IMG-2...).
    Deduplica por conteúdo do arquivo (hash) para evitar imagens visualmente iguais."""
    import hashlib
    catalog = []
    seen_paths: set[str] = set()
    seen_hashes: set[str] = set()

    def _add(img_path: str, context: str, source: str) -> None:
        if img_path in seen_paths:
            return
        seen_paths.add(img_path)
        p = Path(img_path)
        if not p.exists():
            return
        # Deduplica por conteúdo
        try:
            content_hash = hashlib.md5(p.read_bytes(), usedforsecurity=False).hexdigest()
        except Exception:
            content_hash = img_path
        if content_hash in seen_hashes:
            return
        seen_hashes.add(content_hash)
        catalog.append({"path": str(p), "context": context, "source": source})

    for block in bundle.market_ear_blocks:
        # Contexto rico: título + subtítulo + primeiros ~350 chars do body
        _title = (block.title or "").strip()
        _sub = (block.subtitle or "").strip()
        _body = (block.body_text or "").strip()
        parts = []
        if _title:
            parts.append(_title)
        if _sub and _sub != _title:
            parts.append(_sub)
        if _body:
            parts.append(_body[:350])
        ctx = " — ".join(parts)
        for img_path in (block.image_refs or []):
            # Adiciona filename como hint adicional (às vezes descritivo)
            fname = Path(img_path).stem
            ctx_full = f"{ctx} [file:{fname}]" if fname else ctx
            _add(img_path, ctx_full, "ZeroHedge")

    for item in bundle.x_items:
        _author = (item.author or "").strip()
        _text = (item.text or "").strip()
        ctx = f"@{_author}: {_text[:350]}" if _author else _text[:350]
        for img_path in (item.media_refs or []):
            _add(img_path, ctx, f"@{item.author}")

    for i, img in enumerate(catalog, 1):
        img["id"] = f"IMG-{i}"

    return catalog


def _image_briefing(catalog: list[dict]) -> str:
    """Gera o bloco de briefing de imagens para incluir no prompt de escrita."""
    if not catalog:
        return ""
    subset = catalog[:30]  # máximo 30 para não explodir o contexto
    lines = [
        "\n--- IMAGENS DISPONÍVEIS ---",
        "Lista de imagens reais que você pode inserir no texto. Cada uma tem um contexto",
        "(o bloco/post de onde veio), que descreve aproximadamente o que a imagem mostra.",
        "",
        "⚠️ REGRA ABSOLUTA DE COERÊNCIA:",
        "  - SÓ use uma imagem se o contexto dela bater DIRETAMENTE com o parágrafo onde ela aparece.",
        "  - É MUITO MELHOR escrever o texto sem imagem do que forçar uma imagem genérica.",
        "  - Se nenhuma imagem da lista bater com o argumento do parágrafo, NÃO use imagem ali.",
        "  - A legenda precisa explicar O QUE a imagem mostra E por que está ali naquele ponto.",
        "",
        "Use entre 2 e 5 imagens no texto total (pode ser menos se nenhuma bater).",
        "NUNCA invente um ID. NUNCA repita uma mesma imagem.",
        "",
    ]
    for img in subset:
        lines.append(f"  {img['id']} [{img['source']}]: {img['context']}")
    lines.append(
        "\nFORMATO — insira a marcação DENTRO do texto onde a imagem faz sentido:\n"
        "  Parágrafo argumentando X.\n"
        "  [IMAGEM: IMG-N | legenda que conecta a imagem ao argumento X]\n"
        "  Parágrafo que continua o raciocínio a partir do que a imagem mostra.\n"
        "  ...\n"
        "Regras:\n"
        "  - Use apenas IDs exatos da lista; cada ID NO MÁXIMO UMA VEZ\n"
        "  - Se você tem dúvida se uma imagem encaixa, NÃO use\n"
        "  - Nunca coloque imagem no início ou fim absoluto do texto — sempre entre parágrafos"
    )
    return "\n".join(lines)


def _resolve_image_ids(text: str, catalog: list[dict]) -> str:
    """Substitui [IMAGEM: IMG-N | legenda] por <<<IMG:path|legenda>>> usando lookup direto.
    Cada ID pode ser usado no máximo uma vez (sem duplicatas)."""
    id_to_path = {img["id"]: img["path"] for img in catalog}
    used_ids: set[str] = set()

    def replace(m: re.Match) -> str:
        raw = m.group(1)
        parts = raw.split("|", 1)
        img_id = parts[0].strip()
        caption = parts[1].strip() if len(parts) > 1 else ""

        if img_id in used_ids:
            return ""  # duplicata — remove silenciosamente

        path = id_to_path.get(img_id)
        if path:
            used_ids.add(img_id)
            if caption:
                return f"<<<IMG:{path}|{caption}>>>"
            return f"<<<IMG:{path}>>>"
        return ""

    # Suporta tanto [IMAGEM: IMG-N] quanto [IMAGEM: IMG-N | legenda]
    text = re.sub(r'\[IMAGEM:\s*(IMG-\d+[^\]]*)\]', replace, text)
    # Remove qualquer marcação restante com formato inválido
    text = re.sub(r'\[IMAGEM:[^\]]*\]', '', text)
    embedded = len(used_ids)
    _log.info("image_embed_done", catalog_size=len(catalog), embedded=embedded)
    return text


# ── Exportação para Word ────────────────────────────────────────────────────────

_IMG_SENTINEL_RE = re.compile(r'<<<IMG:(.+?)>>>')


def to_docx(text: str, mode: str, run_date: str) -> bytes:
    """Converte o texto do writer (com sentinels <<<IMG:path>>>) em bytes de .docx."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    def _add_heading(text_: str, level: int) -> None:
        p = doc.add_heading(text_, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_paragraph(text_: str) -> None:
        if not text_.strip():
            return
        p = doc.add_paragraph(text_.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_image(sentinel_content: str) -> None:
        """sentinel_content é 'path' ou 'path|legenda'."""
        parts_ = sentinel_content.split("|", 1)
        img_path = parts_[0].strip()
        caption = parts_[1].strip() if len(parts_) > 1 else ""
        p = Path(img_path)
        if p.exists():
            try:
                doc.add_picture(str(p), width=Inches(5.5))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap_p = doc.add_paragraph(caption)
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap_p.runs:
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            except Exception:
                pass  # imagem corrompida ou formato não suportado

    # Divide o texto em partes: texto ou <<<IMG:...>>>
    parts = _IMG_SENTINEL_RE.split(text)
    # split alterna: [text, sentinel_content, text, sentinel_content, ...]

    for i, part in enumerate(parts):
        if i % 2 == 1:
            # Parte ímpar = sentinel_content (path ou path|caption)
            _add_image(part)
        else:
            # Parte par = texto markdown simples
            for block in part.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("### "):
                    _add_heading(block[4:], 3)
                elif block.startswith("## "):
                    _add_heading(block[3:], 2)
                elif block.startswith("# "):
                    _add_heading(block[2:], 1)
                elif block == "---":
                    doc.add_paragraph("—" * 30)
                else:
                    _add_paragraph(block)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Interface pública ──────────────────────────────────────────────────────────

class WriterOutput:
    def __init__(self, mode: str, focus: str, angle: str, text: str):
        self.mode = mode
        self.focus = focus
        self.angle = angle
        self.text = text
        self.written_at = datetime.now(timezone.utc)

    def __str__(self) -> str:
        return (
            f"[Modo: {self.mode}]\n"
            f"[Foco: {self.focus}]\n"
            f"[Ângulo: {self.angle}]\n\n"
            f"{self.text}"
        )


def write(result: CurationResult, bundle: DailyIngestionBundle | None = None,
          tema_hint: str | None = None,
          mode_override: str | None = None) -> WriterOutput:
    """
    Pipeline de escrita.
    Modo determinado pelo dia da semana (calendário editorial).
    Ângulo e foco determinados pelo conteúdo do dia.
    tema_hint: tema forçado (ex: "After Hormuz — novo regime energético")
    mode_override: força modo editorial (ex: "week_ahead")
    """
    mode = mode_override or _mode_for_date(result.run_date)
    _log.info("writer_start", run_id=result.run_id, run_date=result.run_date,
              mode=mode, has_bundle=bundle is not None)

    angle = _find_angle(result, mode, bundle, tema_hint=tema_hint)
    _log.info("writer_angle_chosen", mode=mode, focus=angle["focus"][:100])

    text = _write(angle, mode, result, bundle)

    # Persiste ângulo + título para evitar repetição nas próximas execuções
    try:
        title = _extract_title_from_text(text)
        _save_angle_log(mode, angle["focus"], angle["angle"],
                        result.run_date or str(date.today()), title=title)
    except Exception as exc:
        _log.warning("angle_log_save_error", error=str(exc))

    return WriterOutput(
        mode=mode,
        focus=angle["focus"],
        angle=angle["angle"],
        text=text,
    )
